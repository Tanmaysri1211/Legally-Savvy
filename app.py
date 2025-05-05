import streamlit as st
from docx import Document
from docx.shared import Pt
import pdfplumber
from xhtml2pdf import pisa
import io
from typing import Tuple

st.set_page_config(page_title="Accessible Legal Toolkit", layout="wide")

# Sidebar
st.sidebar.title("Settings")
tool = st.sidebar.selectbox(
    label="Select a Tool",
    options=[
        "Legal Research Assistant",
        "Case Brief Helper",
        "Legal Drafting Assistant",
        "Contract Drafting Assistant",
        "Contract Analyzer"
    ],
    help="Use arrow keys to navigate between tools"
)
reduce_motion = st.sidebar.checkbox("Reduce Motion", help="Enable this to reduce visual motion for accessibility.")
dark_mode = st.sidebar.checkbox("Enable Dark Mode", help="Use a high-contrast dark theme.")

# Skip link & Styling
st.markdown("""
<a href="#main" class="skip-link" style="position:absolute;top:-40px;left:0;
background:#000;color:#fff;padding:8px;z-index:1000;">Skip to main content</a>
""", unsafe_allow_html=True)

if dark_mode:
    st.markdown("""
    <style>
    html, body, [class*="css"] {
        background-color: #121212 !important;
        color: #FFFFFF !important;
    }
    .stTextInput > div > input, .stTextArea textarea {
        background-color: #1e1e1e !important;
        color: #fff !important;
        border: 1px solid #888 !important;
    }
    .stButton>button {
        background-color: #333 !important;
        color: #fff !important;
        border: 1px solid #fff !important;
    }
    *:focus {
        outline: 3px solid #ffbf47 !important;
    }
    </style>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <style>
    *:focus {
        outline: 3px solid #ffbf47 !important;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown('<h1 id="main" role="heading" aria-level="1">Accessible Legal Toolkit</h1>', unsafe_allow_html=True)

# Utility Functions
def read_file(uploaded_file) -> str:
    file_type = uploaded_file.name.split('.')[-1].lower()
    if file_type == "pdf":
        with pdfplumber.open(uploaded_file) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    if file_type == "docx":
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    if file_type == "txt":
        return uploaded_file.read().decode("utf-8")
    return ""

def save_docx(text: str, filename: str) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    for line in text.split("\n"):
        doc.add_paragraph(line, style='Normal')
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def save_pdf(text: str, filename: str) -> bytes:
    formatted = text.replace("\n", "<br/>")
    html = f"<html><body><h1>{filename}</h1><p>{formatted}</p></body></html>"
    buf = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html), dest=buf)
    return buf.getvalue()

def accessible_download(text: str, name: str):
    pdf_data = save_pdf(text, name)
    docx_data = save_docx(text, name)
    st.download_button("Download as Accessible PDF", data=pdf_data, file_name=f"{name}.pdf", mime="application/pdf")
    st.download_button("Download as Accessible DOCX", data=docx_data, file_name=f"{name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def file_input_area() -> Tuple[str, str]:
    uploaded_file = st.file_uploader("Upload a legal document (.pdf, .docx, .txt)", type=["pdf","docx","txt"])
    text = st.text_area("Or paste your legal text here:", height=300)
    if uploaded_file:
        return read_file(uploaded_file), uploaded_file.name
    return text, "input"

def action_buttons() -> Tuple[bool,bool]:
    c1, c2 = st.columns(2)
    with c1:
        submit = st.button("Submit", help="Click to process")
    with c2:
        cancel = st.button("Cancel", help="Click to cancel")
    return submit, cancel

# Tool implementations
if tool == "Legal Research Assistant":
    st.header("Legal Research Assistant")
    topic = st.text_input("Enter a legal topic:")
    submit, cancel = action_buttons()
    if submit and topic:
        summary = f"Research summary for {topic}"
        st.write(summary)
        accessible_download(summary, "research_summary")

elif tool == "Case Brief Helper":
    st.header("Case Brief Helper")
    txt, name = file_input_area()
    submit, cancel = action_buttons()
    if submit and txt:
        brief = "Facts: ...\nIssues: ...\nArguments: ...\nJudgment: ...\nRatio: ..."
        st.write(brief)
        accessible_download(brief, "case_brief")

elif tool == "Legal Drafting Assistant":
    st.header("Legal Drafting Assistant")
    scenario = st.text_area("Describe scenario:")
    submit, cancel = action_buttons()
    if submit and scenario:
        draft = f"Draft for: {scenario}"
        st.write(draft)
        accessible_download(draft, "draft")

elif tool == "Contract Drafting Assistant":
    st.header("Contract Drafting Assistant")
    ctype = st.text_input("Contract Type:")
    a = st.text_input("Party A:")
    b = st.text_input("Party B:")
    cls = st.text_area("Clauses:")
    submit, cancel = action_buttons()
    if submit and ctype and a and b:
        contract = f"{ctype} between {a} and {b}\nClauses: {cls}"
        st.write(contract)
        accessible_download(contract, "contract")

elif tool == "Contract Analyzer":
    st.header("Contract Analyzer")
    txt, name = file_input_area()
    submit, cancel = action_buttons()
    if submit and txt:
        analysis = "Clause analysis..."
        st.write(analysis)
        accessible_download(analysis, "analysis")
